from docx import Document
from docx.shared import Inches

def createdocument(listoftext, listofimages):
    document = Document()

    print("Attempting to create word document")

    for text, image in zip(listoftext, listofimages):

        try:
            document.add_paragraph(text)
            document.add_picture(image)

        except:
            print("Either text or image not found")
            continue

    print("Finished creating word document")
    document.save(r'C:\Users\Ahmed\Desktop\my programs\PY\Textprocessor\img.docx')




